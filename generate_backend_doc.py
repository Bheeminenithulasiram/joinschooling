"""Script to generate a comprehensive, professionally formatted .docx technical document
detailing the EduConnect / JoinSchooling FastAPI backend architecture and its Next.js frontend integration.
"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # Color Palette
    BRAND_COLOR = RGBColor(79, 70, 229)    # Indigo / Brand
    ACCENT_COLOR = RGBColor(14, 165, 233)  # Sky Blue
    DARK_BG = "1E1B4B"                    # Deep indigo hex
    LIGHT_BG = "F8FAFC"                   # Slate 50 hex
    HEADER_BG = "4F46E5"                  # Brand Indigo hex
    BORDER_COLOR = "E2E8F0"

    # --- TITLE / COVER SECTION ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("EduConnect / JoinSchooling")
    run_title.font.name = 'Segoe UI Semibold'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = BRAND_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Comprehensive Backend Architecture & Frontend API Integration Reference Manual")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(100, 116, 139) # Slate 500

    # Meta Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Architecture Tier", "Decoupled Full-Stack SaaS (FastAPI + Next.js 15 App Router)"),
        ("Database Engine", "Serverless PostgreSQL (Neon.tech) with SQLAlchemy 2.0 ORM"),
        ("Authentication Standard", "JWT (PyJWT) with Argon2 Password Hashing & Silent Cookie/Header Handshake"),
        ("Production Endpoints", "API: https://joinschooling-api-heot.onrender.com | Web: https://joinschooling.vercel.app")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.6)
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)
        
        p_l = cell_lbl.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(9.5)
        
        p_v = cell_val.paragraphs[0]
        r_v = p_v.add_run(val)
        r_v.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    def add_section_header(title, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = BRAND_COLOR
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(15, 23, 42)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.bold = True
            r_b.font.color.rgb = RGBColor(15, 23, 42)
        p.add_run(text)
        return p

    def create_api_card(method, path, auth_req, desc, request_schema, response_schema, fe_usage):
        add_section_header(f"{method}  {path}", level=2)
        
        table = doc.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows_data = [
            ("Authentication", auth_req),
            ("Function & Behavior", desc),
            ("Request Body / Params", request_schema),
            ("Response DTO", response_schema),
            ("Frontend Integration", fe_usage)
        ]
        
        for idx, (label, val) in enumerate(rows_data):
            row = table.rows[idx]
            c_label, c_val = row.cells[0], row.cells[1]
            c_label.width = Inches(1.8)
            c_val.width = Inches(5.0)
            
            set_cell_background(c_label, "F8FAFC")
            set_cell_background(c_val, "FFFFFF")
            set_cell_margins(c_label, 70, 70, 100, 100)
            set_cell_margins(c_val, 70, 70, 100, 100)
            
            p_l = c_label.paragraphs[0]
            r_l = p_l.add_run(label)
            r_l.bold = True
            r_l.font.size = Pt(9.5)
            
            p_v = c_val.paragraphs[0]
            r_v = p_v.add_run(val)
            r_v.font.size = Pt(9.5)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECTION 1: ARCHITECTURE OVERVIEW ---
    add_section_header("1. Executive Summary & Full-Stack Architecture", level=1)
    doc.add_paragraph(
        "EduConnect (JoinSchooling) is engineered as an enterprise-grade education marketplace and career matchmaking platform. "
        "The system separates frontend presentation from core domain business logic through a clean RESTful contract. "
        "The backend is powered by FastAPI (Python 3.12+), utilizing modern asynchronous event-driven design, Pydantic v2 schemas for strict data validation, "
        "SQLAlchemy 2.0 Declarative ORM models, and an automated PostgreSQL database hosted on Neon.tech."
    )
    
    add_bullet(" High-speed ASGI application with automatic OpenAPI (Swagger / ReDoc) generation and typed route parameters.", "FastAPI Framework:")
    add_bullet(" Declarative Base models supporting soft-deletes, automated timestamp mixins, UUID-v4 primary keys, and polymorphic relations.", "SQLAlchemy 2.0 ORM:")
    add_bullet(" Robust JSON Web Token lifecycle via PyJWT with Argon2 hashing algorithm and token blacklisting in PostgreSQL.", "Security & Auth Engine:")
    add_bullet(" Multi-layer middleware handling Request ID propagation, SlowAPI IP/User rate limiting, Strict-Transport-Security (HSTS), and CORS origin parsing.", "Middleware Stack:")
    add_bullet(" Server-Side Rendering (RSC) + Server Actions with an Edge Middleware silent refresh handshake.", "Next.js 15 Client Handshake:")

    # --- SECTION 2: AUTHENTICATION & SESSION LIFECYCLE ---
    add_section_header("2. Authentication & Session Management Lifecycle", level=1)
    doc.add_paragraph(
        "Authentication uses a hardened stateless JWT strategy coupled with a stateful database token revocation mechanism. "
        "The application issues two distinct token types on successful authentication:"
    )
    add_bullet(" Short-lived JWT (15-minute expiration) signed using HS256 algorithm with a high-entropy secret. Carries user ID (sub) and role claim.", "1. Access Token (ec_at):")
    add_bullet(" Long-lived cryptographic token (30-day expiration). Its SHA-256 hash is recorded in the PostgreSQL `refresh_tokens` table. Allows one-time rotation.", "2. Refresh Token (ec_rt):")

    doc.add_paragraph(
        "How the Frontend-Backend Handshake Works:\n"
        "1. When a user logs in via `POST /api/v1/auth/login`, the Next.js Server Action (`loginAction`) securely receives both tokens and sets them as `httpOnly`, `sameSite: lax` cookies in the user's browser.\n"
        "2. For Server Component data fetching (SSR), `frontend/lib/api.ts` extracts `ec_at` and injects `Authorization: Bearer <token>` in the upstream request to FastAPI.\n"
        "3. For Client Components (like `Navbar.tsx`), browser fetch requests hit the Next.js reverse proxy (`/api/v1/:path*`). FastAPI's dependency `get_current_user` inspects the `Authorization` header first, and seamlessly falls back to reading the `ec_at` cookie if the header is absent.\n"
        "4. When an Access Token expires, Next.js Edge Middleware (`middleware.ts`) automatically intercepts the request, calls `POST /api/v1/auth/refresh`, updates the cookies in-flight, and serves the user without any session interruption."
    )

    # --- SECTION 3: DETAILED API ENDPOINTS ---
    add_section_header("3. Comprehensive API Endpoints Specification", level=1)

    # 3.1 Meta & Health Endpoints
    create_api_card(
        method="GET",
        path="/healthz",
        auth_req="Public (No auth required)",
        desc="Kubernetes/Render liveness probe. Returns process uptime and server timestamp.",
        request_schema="None",
        response_schema='{"status": "ok", "uptime_s": 142.5, "ts": "2026-08-31T22:45:00Z"}',
        fe_usage="Used by cloud load balancers and Render health monitors to verify the web process is running."
    )

    create_api_card(
        method="GET",
        path="/readyz",
        auth_req="Public (No auth required)",
        desc="Readiness probe. Executes `SELECT 1` on the PostgreSQL connection to confirm database connectivity.",
        request_schema="None",
        response_schema='{"status": "ready", "db": "ok"} (HTTP 200) or {"status": "not_ready", "db": "error"} (HTTP 503)',
        fe_usage="Used during deployment verification to ensure database migrations and connection pool checkouts succeed."
    )

    # 3.2 Authentication Endpoints
    create_api_card(
        method="POST",
        path="/api/v1/auth/register",
        auth_req="Public | Rate Limit: 3/min, 10/hour",
        desc="Creates a new User record and associated Student profile. Hashes the password using Argon2 and creates initial tokens.",
        request_schema='{\n  "email": "student@educonnect.dev",\n  "password": "securepassword123",\n  "first_name": "Ravi",\n  "last_name": "Kumar",\n  "role": "student"\n}',
        response_schema='{\n  "access_token": "eyJhbGciOi...",\n  "refresh_token": "e89a...",\n  "token_type": "bearer",\n  "expires_in": 900\n} (HTTP 201)',
        fe_usage="Invoked by `registerAction()` in `frontend/app/auth/register/page.tsx`. Automatically stores cookies and redirects to `/dashboard`."
    )

    create_api_card(
        method="POST",
        path="/api/v1/auth/login",
        auth_req="Public | Rate Limit: 5/min, 20/hour",
        desc="Validates user credentials against Argon2 hash. Generates new access and refresh tokens.",
        request_schema='{\n  "email": "student@educonnect.dev",\n  "password": "securepassword123"\n}',
        response_schema='{\n  "access_token": "eyJhbGciOi...",\n  "refresh_token": "e89a...",\n  "token_type": "bearer",\n  "expires_in": 900\n} (HTTP 200)',
        fe_usage="Invoked by `loginAction()` in `frontend/app/auth/login/page.tsx` on credential submission."
    )

    create_api_card(
        method="POST",
        path="/api/v1/auth/refresh",
        auth_req="Public (Requires valid Refresh Token) | Rate Limit: 30/min",
        desc="Rotates refresh token. Hashes input token, checks revocation status, issues a fresh token pair, and marks the old token as revoked.",
        request_schema='{\n  "refresh_token": "e89a..."\n}',
        response_schema='{\n  "access_token": "eyJhbGciOi...",\n  "refresh_token": "f90b...",\n  "token_type": "bearer",\n  "expires_in": 900\n} (HTTP 200)',
        fe_usage="Called automatically by Next.js `middleware.ts` and `api.ts` interceptor when `ec_at` expires."
    )

    create_api_card(
        method="POST",
        path="/api/v1/auth/logout",
        auth_req="Authenticated",
        desc="Revokes the active refresh token in the database, preventing future token rotations.",
        request_schema='{\n  "refresh_token": "e89a..."\n}',
        response_schema='HTTP 204 No Content',
        fe_usage="Triggered by `logoutAction()` in `frontend/components/layout/Navbar.tsx`. Clears browser cookies and redirects to homepage."
    )

    # 3.3 User & Dashboard Endpoints
    create_api_card(
        method="GET",
        path="/api/v1/me",
        auth_req="Bearer Token or `ec_at` Cookie",
        desc="Returns the current authenticated user record with full profile details (CGPA, marks, skills, dream companies).",
        request_schema="None",
        response_schema='{\n  "id": "uuid",\n  "email": "student@educonnect.dev",\n  "role": "student",\n  "is_email_verified": true,\n  "profile": {\n    "first_name": "Ravi",\n    "last_name": "Kumar",\n    "cgpa": 8.6,\n    "preferred_course": "Computer Science Engineering",\n    "skills": ["Python", "React"]\n  }\n}',
        fe_usage="Used by `Navbar.tsx` on client mount to dynamically toggle between guest and logged-in states, and by `/dashboard`."
    )

    create_api_card(
        method="GET",
        path="/api/v1/me/dashboard",
        auth_req="Bearer Token or `ec_at` Cookie",
        desc="Aggregates student overview: application count, saved bookmarks count, unread notifications, recent applications, and recommended colleges.",
        request_schema="None",
        response_schema='{\n  "stats": {"applications": 3, "saved_colleges": 5, "saved_internships": 2, "unread_notifs": 1},\n  "recent_applications": [...],\n  "recommended_colleges": [...],\n  "upcoming_deadlines": []\n}',
        fe_usage="Loaded server-side by `frontend/app/dashboard/page.tsx` for real-time student analytics."
    )

    create_api_card(
        method="GET",
        path="/api/v1/me/notifications",
        auth_req="Bearer Token or `ec_at` Cookie (Student Role)",
        desc="Retrieves the chronological stream of in-app notifications (application status changes, deadlines, AI match alerts).",
        request_schema="None",
        response_schema='[\n  {\n    "id": "uuid",\n    "kind": "application",\n    "title": "Amazon SDE moved to Under Review",\n    "body": "Recruiter viewed your profile",\n    "href": "/applications",\n    "read": false,\n    "created_at": "2026-08-31T18:00:00Z"\n  }\n]',
        fe_usage="Rendered by `frontend/app/notifications/page.tsx` with category badges and direct navigation links."
    )

    create_api_card(
        method="POST",
        path="/api/v1/me/notifications/read",
        auth_req="Bearer Token or `ec_at` Cookie",
        desc="Marks all unread notifications as read by updating `read_at` timestamp in PostgreSQL.",
        request_schema="None",
        response_schema='HTTP 204 No Content',
        fe_usage="Invoked by Server Action `markAllNotificationsReadAction()` when user clicks 'Mark all read' button."
    )

    # 3.4 Colleges Endpoints
    create_api_card(
        method="GET",
        path="/api/v1/colleges",
        auth_req="Public",
        desc="Paginated search and discovery for colleges. Supports filtering by name query (q), state, type (private, government, deemed), and multi-field sorting.",
        request_schema="Query params: `q`, `state`, `type`, `sort` (rating, avg_package, nirf_rank, fees_asc), `page`, `page_size`",
        response_schema='{\n  "items": [\n    {\n      "id": "uuid",\n      "slug": "iit-bombay",\n      "name": "IIT Bombay",\n      "banner_url": "https://images.unsplash.com/...",\n      "nirf_rank": 3,\n      "avg_package_lpa": 21.5,\n      "placement_percent": 96.0,\n      "rating": 4.9\n    }\n  ],\n  "pagination": {"page": 1, "page_size": 24, "total": 5, "has_next": false}\n}',
        fe_usage="Powers the directory view in `frontend/app/colleges/page.tsx` with dynamic URL search parameters and responsive card grids."
    )

    create_api_card(
        method="GET",
        path="/api/v1/colleges/{slug}",
        auth_req="Public",
        desc="Fetches rich detail view for a specific college by unique URL slug. Includes facilities, courses, fee tiers, and placement trends.",
        request_schema="Path param: `slug` (e.g. `vnr-vjiet`, `iit-bombay`)",
        response_schema='{\n  "id": "uuid",\n  "slug": "vnr-vjiet",\n  "name": "VNR VJIET",\n  "about": "Autonomous institute...",\n  "banner_url": "https://...",\n  "facilities": ["Library", "Hostel", "Sports"],\n  "courses": [{"name": "CSE", "duration_years": 4, "fees_per_year_lpa": 1.35}],\n  "hostel_available": true\n}',
        fe_usage="Rendered by `frontend/app/colleges/[slug]/page.tsx` with full-bleed hero banner, courses breakdown, and interactive bookmarking."
    )

    create_api_card(
        method="GET",
        path="/api/v1/colleges/compare",
        auth_req="Public",
        desc="Returns multi-college comparison records matching provided list of slugs or IDs.",
        request_schema="Query param: `ids` (comma-separated IDs or slugs)",
        response_schema='[ CollegeDetail, CollegeDetail, ... ]',
        fe_usage="Powers the dynamic 4-column comparison table on `frontend/app/compare/page.tsx`."
    )

    # 3.5 Internships Endpoints
    create_api_card(
        method="GET",
        path="/api/v1/internships",
        auth_req="Public",
        desc="Paginated internship search. Supports filtering by domain (Software, Data, Design), work_mode (remote, hybrid, onsite), and min_stipend.",
        request_schema="Query params: `q`, `domain`, `work_mode`, `min_stipend`, `page`, `page_size`",
        response_schema='{\n  "items": [\n    {\n      "id": "uuid",\n      "slug": "software-development-intern-amazon",\n      "title": "Software Development Intern",\n      "domain": "Software",\n      "work_mode": "remote",\n      "duration_months": 2,\n      "stipend_min": 30000,\n      "stipend_max": 30000,\n      "company": {"name": "Amazon", "logo_url": "https://logo.clearbit.com/amazon.com"}\n    }\n  ],\n  "pagination": {"page": 1, "page_size": 24, "total": 2, "has_next": false}\n}',
        fe_usage="Populates the listing page on `frontend/app/internships/page.tsx` with company logos and stipend summary chips."
    )

    create_api_card(
        method="POST",
        path="/api/v1/internships/{id}/apply",
        auth_req="Authenticated (Student Role)",
        desc="Submits an internship application. Validates that the internship exists, verifies active status, and prevents duplicate submissions (409 Conflict).",
        request_schema='{\n  "cover_letter": "I am proficient in Python and React...",\n  "answers": {"availability": "immediate"}\n}',
        response_schema='{\n  "id": "uuid",\n  "user_id": "uuid",\n  "target_kind": "internship",\n  "target_id": "uuid",\n  "status": "submitted",\n  "submitted_at": "2026-08-31T22:50:00Z"\n} (HTTP 201)',
        fe_usage="Triggered by the `ApplyButton` component on `frontend/app/internships/[slug]/ApplyButton.tsx`."
    )

    # 3.6 Bookmarks & Saved Items
    create_api_card(
        method="POST",
        path="/api/v1/saved",
        auth_req="Authenticated",
        desc="Bookmarks a college or internship. Performs polymorphic database validation to guarantee the target ID exists, preventing orphaned references.",
        request_schema='{\n  "kind": "college",\n  "target_id": "uuid"\n}',
        response_schema='{\n  "id": "uuid",\n  "user_id": "uuid",\n  "kind": "college",\n  "target_id": "uuid",\n  "created_at": "2026-08-31T22:50:00Z"\n} (HTTP 201)',
        fe_usage="Used by the `SaveButton` client component (`frontend/components/ui/SaveButton.tsx`) across college and internship cards."
    )

    create_api_card(
        method="DELETE",
        path="/api/v1/saved/{item_id}",
        auth_req="Authenticated",
        desc="Removes a bookmarked item by ID for the logged-in user.",
        request_schema="Path param: `item_id`",
        response_schema='HTTP 204 No Content',
        fe_usage="Triggered when un-bookmarking items in `/colleges/saved` or `/internships/saved`."
    )

    # 3.7 AI College Finder
    create_api_card(
        method="POST",
        path="/api/v1/ai/college-finder",
        auth_req="Optional Auth (Publicly accessible)",
        desc="Executes the multi-factor weighted college recommendation algorithm. Computes match scores and admission probability.",
        request_schema='{\n  "tenth_percentage": 92,\n  "twelfth_percentage": 88,\n  "cgpa": 8.5,\n  "preferred_course": "Computer Science Engineering",\n  "budget_max_lpa": 3.0,\n  "state": "Telangana",\n  "hostel_required": true,\n  "expected_package_lpa": 12,\n  "preferred_companies": ["Amazon", "Microsoft"]\n}',
        response_schema='{\n  "run_id": "uuid",\n  "recommendations": [\n    {\n      "college": { ... },\n      "match_score": 0.94,\n      "admission_probability": 0.88,\n      "matching_reasons": ["Top placement tier matching expected LPA", "Course offered within budget cap"]\n    }\n  ]\n}',
        fe_usage="Invoked by `AiFinderForm.tsx` in `frontend/app/ai-finder/page.tsx` to generate instantaneous personalized rankings."
    )

    # --- SECTION 4: DATABASE SCHEMA & ORM MODELING ---
    add_section_header("4. Database Schema & Relational Modeling", level=1)
    doc.add_paragraph(
        "The database layer is managed using SQLAlchemy 2.0 with a clean object-relational mapping design. Key architectural features include:"
    )
    add_bullet(" All tables use UUIDv4 strings for primary keys to prevent enumeration attacks and simplify distributed replication.", "1. UUID Primary Keys:")
    add_bullet(" `created_at` and `updated_at` timestamps are automatically populated via server default `func.now()` and on-update triggers.", "2. Timestamp Mixins:")
    add_bullet(" Tables such as `colleges`, `companies`, and `internships` implement a `deleted_at` nullable timestamp column for non-destructive data lifecycle.", "3. Soft Deletion:")
    add_bullet(" `RefreshToken.token_hash` has a unique database index to ensure O(1) query performance during authentication checks.", "4. Indexed Lookups:")
    add_bullet(" Multi-column unique constraints on `(user_id, kind, target_id)` prevent duplicate bookmarks and application conflicts.", "5. Constraint Hardening:")

    # Models Summary Table
    add_section_header("Database Entities Summary", level=2)
    tbl_models = doc.add_table(rows=7, cols=3)
    tbl_models.alignment = WD_TABLE_ALIGNMENT.CENTER
    models_data = [
        ("Table Name", "Primary Attributes", "Relationships & Constraints"),
        ("users", "id, email, password_hash, role, is_active, is_email_verified", "1-to-1 with `students`, 1-to-Many with `refresh_tokens`"),
        ("students", "user_id, first_name, last_name, cgpa, tenth_pct, twelfth_pct, skills", "Foreign Key to `users.id` with CASCADE delete"),
        ("refresh_tokens", "id, user_id, token_hash, expires_at, revoked_at", "Indexed `token_hash`, Foreign Key to `users.id`"),
        ("colleges", "id, slug, name, type, city, state, nirf_rank, avg_package_lpa, banner_url", "1-to-Many with `courses` and `placement_records`"),
        ("internships", "id, company_id, title, slug, domain, work_mode, stipend_min, stipend_max", "Foreign Key to `companies.id`, 1-to-Many with `applications`"),
        ("saved_items", "id, user_id, kind, target_id, created_at", "Unique Constraint on `(user_id, kind, target_id)`")
    ]
    for idx, (col1, col2, col3) in enumerate(models_data):
        row = tbl_models.rows[idx]
        c1, c2, c3 = row.cells[0], row.cells[1], row.cells[2]
        c1.width = Inches(1.3)
        c2.width = Inches(2.7)
        c3.width = Inches(2.8)
        
        is_hdr = (idx == 0)
        set_cell_background(c1, "F1F5F9" if is_hdr else "FFFFFF")
        set_cell_background(c2, "F1F5F9" if is_hdr else "FFFFFF")
        set_cell_background(c3, "F1F5F9" if is_hdr else "FFFFFF")
        set_cell_margins(c1, 60, 60, 80, 80)
        set_cell_margins(c2, 60, 60, 80, 80)
        set_cell_margins(c3, 60, 60, 80, 80)
        
        for c, text in [(c1, col1), (c2, col2), (c3, col3)]:
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.0)
            if is_hdr:
                r.bold = True

    # --- SECTION 5: SECURITY & PRODUCTION HARDENING ---
    add_section_header("5. Production Security & Ops Hardening", level=1)
    doc.add_paragraph(
        "The backend has been configured to meet strict industry-standard SaaS security benchmarks:"
    )
    add_bullet(" Injected via `SecurityHeadersMiddleware`. Protects against XSS, clickjacking (`X-Frame-Options: DENY`), and MIME sniffing (`X-Content-Type-Options: nosniff`).", "Security Headers:")
    add_bullet(" Powered by `slowapi` to mitigate credential-stuffing and automated DDoS on `/auth/login` and `/auth/register`.", "Rate Limiting:")
    add_bullet(" Automatically assigns a unique UUID tracing header (`X-Request-ID`) to every incoming request for structured error diagnostics.", "Request ID Tracing:")
    add_bullet(" Pydantic `SettingsConfigDict` automatically validates environment variables and normalizes database strings.", "Configuration Safety:")

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Save Document
    os.makedirs("docs", exist_ok=True)
    out_path_1 = "docs/EduConnect_Backend_API_Documentation.docx"
    out_path_2 = "Backend_API_Documentation.docx"
    doc.save(out_path_1)
    doc.save(out_path_2)
    print(f"Document successfully generated at {out_path_1} and {out_path_2}")

if __name__ == "__main__":
    create_document()
